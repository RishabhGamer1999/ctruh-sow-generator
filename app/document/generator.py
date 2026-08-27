"""
CTRUH SOW / Scope of Project Document Generator.
Matches the official 'Scope and Commercials - Format (New Letterhead)' design:
  - Header with CTRUH contact info
  - Solid Blue Banner for 'Project Scope & Commercials'
  - 9 Standard Numbered Sections with CTRUH Blue Headings
  - Formatted Tables (Timeline, Commercials, Add-ons)
  - Footer with Company Name, CIN, and Registered Address
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


# ─── CTRUH Brand Colors ───────────────────────────────────────────────────────
COLOR_BLUE       = RGBColor(0x00, 0x66, 0xFF)  # #0066FF (CTRUH primary blue)
COLOR_DARK_TEXT  = RGBColor(0x1F, 0x29, 0x37)  # #1F2937 (Charcoal dark text)
COLOR_MUTED_TEXT = RGBColor(0x6B, 0x72, 0x80)  # #6B7280 (Muted gray)
COLOR_LIGHT_BG   = "F3F4F6"                    # Soft background for table alternating rows
COLOR_BANNER_BG  = "0066FF"                    # CTRUH Blue banner fill
FONT_NAME        = "Arial"                     # Professional clean font


def _set_cell_bg(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)


def _set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _format_run(run, font_name=FONT_NAME, size_pt=10, bold=False, italic=False, color_rgb=COLOR_DARK_TEXT):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color_rgb


def _add_section_heading(doc: Document, number_title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(number_title)
    _format_run(r, size_pt=11.5, bold=True, color_rgb=COLOR_BLUE)
    return p


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    _format_run(r, size_pt=9.5, color_rgb=COLOR_DARK_TEXT)
    return p


def generate_sop_docx(data: dict) -> bytes:
    """
    Generate the official CTRUH Scope & Commercials .docx file.
    """
    doc = Document()

    # Set default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(9.5)
    font.color.rgb = COLOR_DARK_TEXT

    # ── Page Margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.85)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)

    client_name = data.get("client_name", "Client Partner")
    today_str   = data.get("date", datetime.now().strftime("%d %B %Y"))

    # ── Header Contact Info ───────────────────────────────────────────────────
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(8)
    
    r_brand = header_p.add_run("CTRUH\n")
    _format_run(r_brand, size_pt=18, bold=True, color_rgb=COLOR_BLUE)

    r_contact = header_p.add_run("08047363099  |  hello@ctruh.com  |  www.ctruh.com")
    _format_run(r_contact, size_pt=8.5, color_rgb=COLOR_MUTED_TEXT)

    # ── Solid Blue Title Banner ───────────────────────────────────────────────
    banner_tbl = doc.add_table(rows=1, cols=1)
    banner_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner_tbl.autofit = False

    b_cell = banner_tbl.rows[0].cells[0]
    b_cell.width = Inches(6.8)
    _set_cell_bg(b_cell, COLOR_BANNER_BG)
    _set_cell_margins(b_cell, top=140, bottom=140, left=180, right=180)

    bp = b_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_btitle = bp.add_run("Project Scope & Commercials")
    _format_run(r_btitle, size_pt=15, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))

    # ── Subtitle Info (Prepared for / Date) ───────────────────────────────────
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(6)
    sub_p.paragraph_format.space_after = Pt(12)

    r_prep = sub_p.add_run("Prepared for: ")
    _format_run(r_prep, size_pt=9.5, color_rgb=COLOR_MUTED_TEXT)

    r_cname = sub_p.add_run(f"{client_name}")
    _format_run(r_cname, size_pt=9.5, bold=True, color_rgb=COLOR_DARK_TEXT)

    r_sep = sub_p.add_run(f"    |    Date: {today_str}")
    _format_run(r_sep, size_pt=9.5, color_rgb=COLOR_MUTED_TEXT)

    # ── 1. Project Objective ──────────────────────────────────────────────────
    _add_section_heading(doc, "1. Project Objective")
    obj_p = doc.add_paragraph()
    obj_p.paragraph_format.space_after = Pt(6)
    obj_p.paragraph_format.line_spacing = 1.15
    obj_text = data.get("project_objective") or data.get("project_description") or (
        f"To deliver digital experiences and marketing creatives for {client_name}, "
        f"based on agreed requirements, assets, and timelines."
    )
    r_obj = obj_p.add_run(obj_text)
    _format_run(r_obj, size_pt=9.5)

    # ── 2. Scope of Work ──────────────────────────────────────────────────────
    _add_section_heading(doc, "2. Scope of Work")
    intro_p = doc.add_paragraph("Ctruh will deliver the following:")
    intro_p.paragraph_format.space_after = Pt(3)
    _format_run(intro_p.runs[0], size_pt=9.5)

    in_scope = data.get("in_scope", [])
    if isinstance(in_scope, list) and in_scope:
        for item in in_scope:
            _add_bullet(doc, item)
    else:
        _add_bullet(doc, str(in_scope or "Deliverables as specified in requirements."))

    # ── 3. Timeline ───────────────────────────────────────────────────────────
    _add_section_heading(doc, "3. Timeline")
    
    phases = data.get("timeline_phases", [
        {"phase": "Requirement & asset collection", "timeline": "3-5 Working Days"},
        {"phase": "First version / preview",        "timeline": "2 Weeks"},
        {"phase": "Feedback implementation",       "timeline": "1 Week"},
        {"phase": "Final delivery",                "timeline": "Within 4-5 Weeks"},
    ])

    tbl = doc.add_table(rows=len(phases) + 1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = "Phase"
    hdr_cells[1].text = "Timeline"
    hdr_cells[0].width = Inches(4.3)
    hdr_cells[1].width = Inches(2.5)

    for cell in hdr_cells:
        _set_cell_bg(cell, COLOR_BANNER_BG)
        _set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
        p = cell.paragraphs[0]
        if p.runs:
            _format_run(p.runs[0], size_pt=9.5, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))

    for i, phase_item in enumerate(phases):
        row_cells = tbl.rows[i + 1].cells
        row_cells[0].width = Inches(4.3)
        row_cells[1].width = Inches(2.5)

        if isinstance(phase_item, dict):
            row_cells[0].text = phase_item.get("phase", "")
            row_cells[1].text = phase_item.get("timeline", "")
        else:
            row_cells[0].text = str(phase_item)
            row_cells[1].text = "—"

        for cell in row_cells:
            bg = COLOR_LIGHT_BG if (i % 2 == 1) else "FFFFFF"
            _set_cell_bg(cell, bg)
            _set_cell_margins(cell, top=60, bottom=60, left=140, right=140)
            if cell.paragraphs[0].runs:
                _format_run(cell.paragraphs[0].runs[0], size_pt=9)

    tl_note = doc.add_paragraph()
    tl_note.paragraph_format.space_before = Pt(4)
    tl_note.paragraph_format.space_after = Pt(6)
    est_str = data.get("timeline_estimate", "Estimated delivery: based on scope and asset approvals.")
    r_n1 = tl_note.add_run(f"{est_str}\n")
    _format_run(r_n1, size_pt=8.5, italic=True, color_rgb=COLOR_MUTED_TEXT)
    r_n2 = tl_note.add_run("Timelines are subject to timely client inputs and feedback.")
    _format_run(r_n2, size_pt=8.5, italic=True, color_rgb=COLOR_MUTED_TEXT)

    # ── 4. Iterations Included ────────────────────────────────────────────────
    _add_section_heading(doc, "4. Iterations Included")
    rounds = data.get("iteration_rounds", "2")
    it_p = doc.add_paragraph(f"This scope includes {rounds} rounds of revisions.")
    it_p.paragraph_format.space_after = Pt(3)
    _format_run(it_p.runs[0], size_pt=9.5)

    it_items = data.get("iteration_items", [
        "Text corrections and content updates",
        "Minor animation adjustments",
        "Timing and transition refinements",
        "Brand alignment and styling checks",
    ])
    for item in it_items:
        _add_bullet(doc, item)

    it_note = doc.add_paragraph("Major creative direction changes or additional concepts will be treated as change requests and quoted separately.")
    it_note.paragraph_format.space_before = Pt(3)
    it_note.paragraph_format.space_after = Pt(6)
    _format_run(it_note.runs[0], size_pt=8.5, italic=True, color_rgb=COLOR_MUTED_TEXT)

    # ── 5. Out of Scope ───────────────────────────────────────────────────────
    _add_section_heading(doc, "5. Out of Scope")
    oos_intro = doc.add_paragraph("The following are not included unless agreed separately:")
    oos_intro.paragraph_format.space_after = Pt(3)
    _format_run(oos_intro.runs[0], size_pt=9.5)

    out_of_scope = data.get("out_of_scope", [
        "Product photography",
        "3D modelling from scratch unless specified",
        "Creation of additional product assets outside scope",
        "Third-party integrations",
        "Extra revision rounds beyond agreed limit",
    ])
    for item in out_of_scope:
        _add_bullet(doc, item)

    # ── 6. Acceptance Criteria ────────────────────────────────────────────────
    _add_section_heading(doc, "6. Acceptance Criteria")
    ac_intro = doc.add_paragraph("A deliverable is considered complete and accepted when:")
    ac_intro.paragraph_format.space_after = Pt(3)
    _format_run(ac_intro.runs[0], size_pt=9.5)

    ac_items = data.get("acceptance_criteria", [
        "All items in the Scope of Work are delivered and accessible.",
        "Everything works as specified, with no critical errors.",
        "The output matches the approved designs, references, and brand guidelines.",
        "Deliverables are accessible via agreed format or platform link.",
        "Client confirms acceptance in writing.",
    ])
    for item in ac_items:
        _add_bullet(doc, item)

    ac_clause = doc.add_paragraph("If no feedback is received within 5 business days of delivery, the deliverables shall be considered approved.")
    ac_clause.paragraph_format.space_before = Pt(3)
    ac_clause.paragraph_format.space_after = Pt(6)
    _format_run(ac_clause.runs[0], size_pt=8.5, italic=True, color_rgb=COLOR_MUTED_TEXT)

    # ── 7. Commercials ────────────────────────────────────────────────────────
    _add_section_heading(doc, "7. Commercials")
    comm_intro = doc.add_paragraph("The plan below lists the full pricing and deliverables:")
    comm_intro.paragraph_format.space_after = Pt(6)
    _format_run(comm_intro.runs[0], size_pt=9.5)

    packages = data.get("commercial_packages", [])
    if not packages and data.get("pricing"):
        packages = [{
            "title": "The Plan",
            "subtitle": f"For {client_name}: Deliverables & Scope",
            "price": str(data.get("pricing", "₹ On Request")),
            "features": data.get("in_scope", [])[:4]
        }]

    for pkg in packages:
        pkg_tbl = doc.add_table(rows=3, cols=1)
        pkg_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        pkg_tbl.autofit = False

        c0 = pkg_tbl.rows[0].cells[0]
        c0.width = Inches(6.8)
        _set_cell_bg(c0, COLOR_BANNER_BG)
        _set_cell_margins(c0, top=70, bottom=70, left=140, right=140)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(pkg.get("title", "The Plan"))
        _format_run(r0, size_pt=10.5, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
        if pkg.get("subtitle"):
            p0.add_run(f"\n{pkg['subtitle']}")
            if len(p0.runs) > 1:
                _format_run(p0.runs[1], size_pt=8.5, color_rgb=RGBColor(0xEE, 0xEE, 0xEE))

        c1 = pkg_tbl.rows[1].cells[0]
        c1.width = Inches(6.8)
        _set_cell_bg(c1, "EBF3FF")
        _set_cell_margins(c1, top=80, bottom=80, left=140, right=140)
        p1 = c1.paragraphs[0]
        r_pr = p1.add_run(pkg.get("price", ""))
        _format_run(r_pr, size_pt=13.5, bold=True, color_rgb=COLOR_BLUE)

        c2 = pkg_tbl.rows[2].cells[0]
        c2.width = Inches(6.8)
        _set_cell_bg(c2, "FAFAFA")
        _set_cell_margins(c2, top=80, bottom=80, left=140, right=140)
        features = pkg.get("features", [])
        if features:
            p2 = c2.paragraphs[0]
            for f_idx, feat in enumerate(features):
                if f_idx > 0:
                    p2 = c2.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                r_feat = p2.add_run(f"✓  {feat}")
                _format_run(r_feat, size_pt=9)

        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # Optional Add-on table
    addons = data.get("addons", [])
    if addons:
        add_p = doc.add_paragraph("Infrastructure & Ongoing Cost")
        _format_run(add_p.runs[0], size_pt=9.5, bold=True)
        
        add_tbl = doc.add_table(rows=len(addons) + 1, cols=3)
        add_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        a_hdrs = add_tbl.rows[0].cells
        a_hdrs[0].text = "Add-on"
        a_hdrs[1].text = "What's included"
        a_hdrs[2].text = "Price / mo"
        a_hdrs[0].width = Inches(2.2)
        a_hdrs[1].width = Inches(3.1)
        a_hdrs[2].width = Inches(1.5)

        for c in a_hdrs:
            _set_cell_bg(c, COLOR_BANNER_BG)
            _set_cell_margins(c, top=60, bottom=60, left=100, right=100)
            if c.paragraphs[0].runs:
                _format_run(c.paragraphs[0].runs[0], size_pt=9, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))

        for a_i, add_item in enumerate(addons):
            r_cells = add_tbl.rows[a_i + 1].cells
            r_cells[0].text = add_item.get("addon", "")
            r_cells[1].text = add_item.get("included", "")
            r_cells[2].text = add_item.get("price", "")
            for c in r_cells:
                _set_cell_bg(c, COLOR_LIGHT_BG if (a_i % 2 == 1) else "FFFFFF")
                _set_cell_margins(c, top=50, bottom=50, left=100, right=100)
                if c.paragraphs[0].runs:
                    _format_run(c.paragraphs[0].runs[0], size_pt=8.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # Billing terms
    bill_head = doc.add_paragraph("Billing & Renewal Terms")
    _format_run(bill_head.runs[0], size_pt=9.5, bold=True)
    bill_head.paragraph_format.space_after = Pt(2)

    billing_terms = data.get("billing_terms", [
        "Project Kickoff Advance - 70%",
        "2 Weeks Post Project Kickoff - 20%",
        "After The Final Delivery - 10%",
        "GST applicable as per government regulations.",
        "Proposal validity: 30 days from the issue date.",
    ])
    for b_term in billing_terms:
        _add_bullet(doc, b_term)

    # ── 8. Client Inputs Required ─────────────────────────────────────────────
    _add_section_heading(doc, "8. Client Inputs Required")
    ci_p = doc.add_paragraph("The client shall provide:")
    ci_p.paragraph_format.space_after = Pt(3)
    _format_run(ci_p.runs[0], size_pt=9.5)

    client_inputs = data.get("client_inputs", [
        "Product images, 3D source files, or specification details",
        "Brand guidelines, fonts, and high-resolution logo files",
        "Preferred references and creative direction guidelines",
        "Timely approvals and consolidated feedback",
    ])
    for c_input in client_inputs:
        _add_bullet(doc, c_input)

    ci_note = doc.add_paragraph("Any delay in inputs or approvals may impact the delivery schedule.")
    ci_note.paragraph_format.space_before = Pt(3)
    ci_note.paragraph_format.space_after = Pt(6)
    _format_run(ci_note.runs[0], size_pt=8.5, italic=True, color_rgb=COLOR_MUTED_TEXT)

    # ── 9. Approval ───────────────────────────────────────────────────────────
    _add_section_heading(doc, "9. Approval")
    appr_p = doc.add_paragraph(
        "Upon written approval, this document shall serve as the agreed Scope of Work, "
        "Commercial Proposal, Delivery Timeline, and Terms of Engagement."
    )
    appr_p.paragraph_format.space_after = Pt(14)
    _format_run(appr_p.runs[0], size_pt=9.5)

    sign_tbl = doc.add_table(rows=2, cols=2)
    sign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_tbl.autofit = False

    for row in sign_tbl.rows:
        for c in row.cells:
            c.width = Inches(3.4)
            _set_cell_margins(c, top=80, bottom=80, left=100, right=100)

    p_client = sign_tbl.rows[0].cells[0].paragraphs[0]
    r_cl = p_client.add_run("Client Approval:\n\n________________________\nDate: _______________")
    _format_run(r_cl, size_pt=9)

    p_ctruh = sign_tbl.rows[0].cells[1].paragraphs[0]
    r_ct = p_ctruh.add_run("Ctruh Stakeholder Approval:\n\n________________________\nDate: _______________")
    _format_run(r_ct, size_pt=9)

    # ── Page Footer (Embedded in document section) ───────────────────────────
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_f1 = footer_p.add_run("CTRUH TECHNOLOGIES PRIVATE LIMITED  |  CIN: U72900KA2022PTC169158\n")
    _format_run(r_f1, size_pt=8, bold=True, color_rgb=COLOR_BLUE)

    r_f2 = footer_p.add_run("3rd Floor, Obeya Silk, 410/381/290, Siddappa Layout, Bommanahalli, Bengaluru - 560068, Karnataka, India")
    _format_run(r_f2, size_pt=7.5, color_rgb=COLOR_MUTED_TEXT)

    # ── Serialize to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
